import os

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Mm

from constants.paths import PATHS

DOORS_LOGO_PATH = os.path.join(PATHS["image"], "doors_logo.jpg")
REPORT_TITLE_IMAGE_PATH = os.path.join(PATHS["image"], "report_title.jpg")

def make_header(
    doc:docx.Document,
    owner_name:str,
    property_name:str,
) -> docx.Document:

    # ___________________content___________________
    # title
    doc.add_picture(REPORT_TITLE_IMAGE_PATH, Mm(150))
    doc.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # owner
    owner_paragraph = doc.add_paragraph(f"オーナー名：{owner_name}").runs[0]
    owner_paragraph.underline = True
    owner_paragraph.bold = True

    # logo
    doc.add_picture(DOORS_LOGO_PATH, Mm(90))
    doc.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # property
    doc.add_table(rows=1, cols=1, style="Table Grid")
    doc.tables[-1].rows[0].cells[0].text = f"物件名：{property_name}"

    # space
    doc.add_paragraph()

    return doc