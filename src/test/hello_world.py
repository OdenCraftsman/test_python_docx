import os

import docx
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

if __name__ == "__main__":
    file_name = os.path.splitext(__file__.split("\\")[-1])[0]

    doc = docx.Document()

    doc.add_paragraph("Hello world!")

    doc.add_paragraph("without underline").runs[0].underline = True
    doc.add_paragraph("bold text").runs[0].bold = True

    doc.add_paragraph("red text").runs[0].font.color.rgb = RGBColor(255, 0, 0)
    doc.add_paragraph("yellow highlight text").runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW

    doc.paragraphs[0].paragraph_format.line_spacing = 1.0
    doc.paragraphs[0].paragraph_format.space_after = 0
    doc.paragraphs[0].paragraph_format.space_before = 0

    doc.save(f"../../test_output/{file_name}.docx")