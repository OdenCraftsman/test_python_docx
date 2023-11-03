import os

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

if __name__ == "__main__":
    file_name = os.path.splitext(__file__.split("\\")[-1])[0]

    doc = docx.Document()

    header = doc.sections[-1].header
    header.add_paragraph("hello header").alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.add_paragraph("hello header").alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.add_paragraph("hello header").alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.add_paragraph("hello header").alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.add_paragraph("hello header").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("hello sections")


    doc.save(f"../../test_output/{file_name}.docx")