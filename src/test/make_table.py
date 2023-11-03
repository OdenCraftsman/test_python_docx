import os

import docx
from docx.shared import RGBColor, Mm
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def set_cell_background_color(cell, color):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    cell._tc.get_or_add_tcPr().append(shading_elm)

if __name__ == "__main__":
    file_name = os.path.splitext(__file__.split("\\")[-1])[0]

    doc = docx.Document()

    doc.add_paragraph("2*2 table")
    table = doc.add_table(rows=2, cols=2, style="Table Grid")
    for cell in table.rows[0].cells:
        set_cell_background_color(cell, "FF0000")


    doc.add_paragraph("1*1 table")
    doc.add_table(rows=1, cols=1, style="Colorful List Accent 1")

    doc.add_paragraph("** *4 table")
    table = doc.add_table(rows=0, cols=0, style="Table Grid")
    table.add_column(Mm(10))
    table.add_column(Mm(40))
    table.add_column(Mm(90))
    table.add_column(Mm(20))

    for i in range(5):
        table.add_row()
    for cell in table.rows[0].cells:
        set_cell_background_color(cell, "C"*6)

    doc.save(f"../../test_output/{file_name}.docx")