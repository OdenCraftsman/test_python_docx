import os

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

if __name__ == "__main__":
    file_name = os.path.splitext(__file__.split("\\")[-1])[0]

    doc = docx.Document()

    doc.add_paragraph("left")
    doc.add_picture("../assets/image/test_image.png", docx.shared.Mm(50))

    doc.add_paragraph("center")
    doc.add_picture("../assets/image/test_image.png", docx.shared.Mm(50))
    doc.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("right")
    doc.add_picture("../assets/image/test_image.png", docx.shared.Mm(50))
    doc.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    para = doc.add_paragraph()
    para.add_run("inline picture")
    para.add_run().add_picture("../assets/image/test_image.png", docx.shared.Mm(50))
    doc.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT


    doc.save(f"../../test_output/{file_name}.docx")